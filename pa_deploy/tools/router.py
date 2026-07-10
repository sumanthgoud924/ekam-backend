import io
import os
import base64
import tempfile
from pathlib import Path
from typing import List, Optional

from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import Response
from pymupdf import open as pdf_open
from PIL import Image

router = APIRouter(prefix="/api/tools", tags=["tools"])

MAX_FILE_SIZE = 50 * 1024 * 1024


@router.post("/image-to-word")
async def image_to_word(file: UploadFile = File(...), language: str = Form("eng")):
    try:
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(413, "File too large. Max 50MB")

        import pytesseract
        from docx import Document
        from docx.shared import Inches

        img = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(img, lang=language)

        doc = Document()
        doc.add_heading("Extracted Text", level=1)
        for paragraph in text.split("\n\n"):
            p = doc.add_paragraph(paragraph.strip())
            if not p.text:
                doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{Path(file.filename).stem}.docx"'},
        )
    except ImportError:
        raise HTTPException(500, "pytesseract or python-docx not installed")
    except Exception as e:
        raise HTTPException(500, str(e))


@router.post("/image-to-pdf")
async def image_to_pdf(files: List[UploadFile] = File(...)):
    try:
        if not files:
            raise HTTPException(400, "No images provided")

        pdf_doc = pdf_open()
        for f in files:
            content = await f.read()
            if len(content) > MAX_FILE_SIZE:
                raise HTTPException(413, f"{f.filename} too large")

            img = Image.open(io.BytesIO(content))
            if img.mode != "RGB":
                img = img.convert("RGB")

            img_bytes = io.BytesIO()
            img.save(img_bytes, format="JPEG", quality=95)
            img_bytes.seek(0)

            img_pdf = pdf_open(img_bytes)
            page = pdf_doc.new_page(width=img_pdf[0].rect.width, height=img_pdf[0].rect.height)
            page.show_pdf_page(page.rect, img_pdf, 0)
            img_pdf.close()

        buf = io.BytesIO()
        pdf_doc.save(buf)
        pdf_doc.close()
        buf.seek(0)

        return Response(
            content=buf.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=\"converted.pdf\""},
        )
    except Exception as e:
        raise HTTPException(500, str(e))


@router.post("/pdf/merge")
async def merge_pdfs(files: List[UploadFile] = File(...)):
    try:
        if len(files) < 2:
            raise HTTPException(400, "Need at least 2 PDF files to merge")

        pdf_writer = pdf_open()
        for f in files:
            content = await f.read()
            if len(content) > MAX_FILE_SIZE:
                raise HTTPException(413, f"{f.filename} too large")
            src = pdf_open(io.BytesIO(content))
            pdf_writer.insert_pdf(src)
            src.close()

        buf = io.BytesIO()
        pdf_writer.save(buf)
        pdf_writer.close()
        buf.seek(0)

        return Response(
            content=buf.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=\"merged.pdf\""},
        )
    except Exception as e:
        raise HTTPException(500, str(e))


@router.post("/pdf/split")
async def split_pdf(
    file: UploadFile = File(...),
    page_ranges: str = Form(...),
):
    try:
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(413, "File too large")

        doc = pdf_open(io.BytesIO(content))
        total_pages = doc.page_count

        ranges = []
        for part in page_ranges.split(","):
            part = part.strip()
            if "-" in part:
                a, b = part.split("-", 1)
                start = int(a.strip()) - 1
                end = int(b.strip())
            else:
                start = int(part) - 1
                end = start + 1
            if start < 0 or end > total_pages:
                doc.close()
                raise HTTPException(400, f"Invalid range '{part}': pages 1-{total_pages}")
            ranges.append((start, end))

        output = pdf_open()
        for start, end in ranges:
            output.insert_pdf(doc, from_page=start, to_page=end - 1)

        doc.close()
        buf = io.BytesIO()
        output.save(buf)
        output.close()
        buf.seek(0)

        return Response(
            content=buf.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=\"split.pdf\""},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


@router.post("/pdf/compress")
async def compress_pdf(file: UploadFile = File(...), quality: int = Form(60)):
    try:
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(413, "File too large")

        quality = max(10, min(95, quality))

        doc = pdf_open(io.BytesIO(content))
        pdf_writer = pdf_open()

        for page_num in range(doc.page_count):
            page = doc[page_num]
            pix = page.get_pixmap(dpi=150)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            img_bytes = io.BytesIO()
            img.save(img_bytes, format="JPEG", quality=quality)
            img_bytes.seek(0)
            img_doc = pdf_open(img_bytes)
            pdf_writer.insert_pdf(img_doc, from_page=0, to_page=0)
            img_doc.close()

        doc.close()
        buf = io.BytesIO()
        pdf_writer.save(buf)
        pdf_writer.close()
        buf.seek(0)

        return Response(
            content=buf.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=\"compressed.pdf\""},
        )
    except Exception as e:
        raise HTTPException(500, str(e))


@router.post("/pdf/rotate")
async def rotate_pdf(file: UploadFile = File(...), angle: int = Form(90)):
    try:
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(413, "File too large")

        angle = angle % 360
        doc = pdf_open(io.BytesIO(content))
        for page in doc:
            page.set_rotation(angle)

        buf = io.BytesIO()
        doc.save(buf)
        doc.close()
        buf.seek(0)

        return Response(
            content=buf.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=\"rotated.pdf\""},
        )
    except Exception as e:
        raise HTTPException(500, str(e))


@router.get("/graphify/codebase")
async def graphify_codebase():
    from tools.graphify_parser import generate_project_graph
    # Resolve the base directory containing the entire repo
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return generate_project_graph(base_dir)

